"""
DOCX Report Generator for Job Market Analytics

This module generates comprehensive written reports in .docx format that combine
all analysis pages into a single, professionally structured document.
"""

import pandas as pd
from typing import Dict, List, Any, Optional
from pathlib import Path
import base64
from io import BytesIO

# Handle optional dependency gracefully
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX report generation will be skipped.")

from .salary_models import SalaryAnalyticsModels
from .nlp_analysis import JobMarketNLPAnalyzer
from .predictive_dashboard import PredictiveAnalyticsDashboard


class JobMarketReportGenerator:
    """
    Generate comprehensive DOCX reports combining all analysis pages.

    Creates professional reports with:
    - Executive summary
    - Model results and explanations
    - Skills analysis and insights
    - Strategic recommendations
    - Appendices with technical details
    """

    def __init__(self, df: pd.DataFrame = None):
        """Initialize report generator with data."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for report generation. "
                "Install with: pip install python-docx"
            )

        # Initialize analytics components using abstraction layer
        self.salary_models = SalaryAnalyticsModels(df)
        self.nlp_analyzer = JobMarketNLPAnalyzer(df)
        self.dashboard = PredictiveAnalyticsDashboard(df)

        # Initialize document
        self.doc = Document()
        self._setup_document_styles()

        print(" DOCX Report Generator initialized")

    def _setup_document_styles(self):
        """Set up document styles for professional formatting."""
        # Title style
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)

        # Heading styles
        heading1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_font = heading1_style.font
        heading1_font.name = 'Arial'
        heading1_font.size = Pt(18)
        heading1_font.bold = True
        heading1_style.paragraph_format.space_before = Pt(18)
        heading1_style.paragraph_format.space_after = Pt(12)

        heading2_style = self.doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_font = heading2_style.font
        heading2_font.name = 'Arial'
        heading2_font.size = Pt(14)
        heading2_font.bold = True
        heading2_style.paragraph_format.space_before = Pt(12)
        heading2_style.paragraph_format.space_after = Pt(6)

        # Body text style
        body_style = self.doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Arial'
        body_font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15

    def generate_complete_report(self, output_path: str = "job_market_analytics_report.docx") -> str:
        """
        Generate complete DOCX report combining all analysis pages.

        Args:
            output_path: Path where to save the DOCX report

        Returns:
            Path to generated report
        """
        print("[DATA] GENERATING COMPREHENSIVE DOCX REPORT")
        print("=" * 50)

        # Run all analyses
        print("[CHECK] Running analytics...")
        analytics_results = self.salary_models.run_complete_analysis()
        nlp_results = self.nlp_analyzer.run_complete_nlp_analysis()
        comprehensive_report = self.dashboard.generate_comprehensive_report()

        # Generate report sections
        self._add_title_page()
        self._add_executive_summary(comprehensive_report)
        self._add_model_1_section(analytics_results['regression'])
        self._add_model_2_section(analytics_results['classification'])
        self._add_nlp_analysis_section(nlp_results)
        self._add_strategic_recommendations(comprehensive_report['recommendations'])
        self._add_technical_appendix(analytics_results, nlp_results)

        # Save document
        output_path = Path(output_path)
        self.doc.save(output_path)

        print(f"[OK] Report generated successfully: {output_path}")
        print(f" Document contains {len(self.doc.paragraphs)} paragraphs")
        print(f" [DATA] File size: {output_path.stat().st_size / 1024:.1f} KB")

        return str(output_path)

    def _add_title_page(self):
        """Add professional title page."""
        # Title
        title = self.doc.add_paragraph("Job Market Analytics Report", style='CustomTitle')

        # Subtitle
        subtitle = self.doc.add_paragraph("Comprehensive Analysis of Salary Trends and Employment Patterns")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_font = subtitle.runs[0].font
        subtitle_font.size = Pt(16)
        subtitle_font.italic = True

        # Add space
        self.doc.add_paragraph()

        # Report details
        details = self.doc.add_paragraph("Report Details:")
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        details_font = details.runs[0].font
        details_font.size = Pt(12)
        details_font.bold = True

        # Analysis components
        components = [
            "• Multiple Linear Regression for Salary Prediction",
            "• Classification Analysis for Above-Average Jobs",
            "• Natural Language Processing for Skills Analysis",
            "• Interactive Dashboards and Visualizations",
            "• Strategic Recommendations for Stakeholders"
        ]

        for component in components:
            comp_para = self.doc.add_paragraph(component)
            comp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            comp_font = comp_para.runs[0].font
            comp_font.size = Pt(11)

        # Add page break
        self.doc.add_page_break()

    def _add_executive_summary(self, comprehensive_report: Dict[str, Any]):
        """Add executive summary section."""
        self.doc.add_paragraph("Executive Summary", style='CustomHeading1')

        exec_summary = comprehensive_report.get('executive_summary', {})

        # Key metrics
        self.doc.add_paragraph("Key Findings", style='CustomHeading2')

        metrics_text = f"""
This comprehensive analysis examined {exec_summary.get('total_jobs_analyzed', 0):,} job postings
to understand salary trends and employment patterns in the current job market.

Our machine learning models achieved the following performance:
• Regression Model R²: {exec_summary.get('regression_r2', 0):.3f} (explains {exec_summary.get('regression_r2', 0)*100:.1f}% of salary variance)
• Classification Accuracy: {exec_summary.get('classification_accuracy', 0):.3f} ({exec_summary.get('classification_accuracy', 0)*100:.1f}% correct predictions)
• Above-Average Salary Threshold: ${exec_summary.get('salary_threshold', 0):,.0f}

The analysis identified key salary drivers and job market patterns that provide actionable
insights for job seekers, employers, and policy makers.
        """

        self.doc.add_paragraph(metrics_text.strip(), style='CustomBody')

        # Top insights
        self.doc.add_paragraph("Primary Insights", style='CustomHeading2')

        insights_text = f"""
Top Salary Driver: {exec_summary.get('top_salary_driver', 'Location and experience level')}
Top Job Predictor: {exec_summary.get('top_job_predictor', 'Industry and skills combination')}

Our analysis reveals significant patterns in compensation that can guide career decisions
and hiring strategies. The models identify specific combinations of location, skills,
and experience that lead to above-average compensation opportunities.
        """

        self.doc.add_paragraph(insights_text.strip(), style='CustomBody')

    def _add_model_1_section(self, regression_results: Dict[str, Any]):
        """Add Model 1: Multiple Linear Regression section."""
        self.doc.add_paragraph("Model 1: Multiple Linear Regression for Salary Prediction", style='CustomHeading1')

        # What we're modeling
        self.doc.add_paragraph("What We're Modeling", style='CustomHeading2')
        modeling_text = """
We're predicting salary based on location, job title, industry, experience, and skills.
This helps understand which factors most influence compensation across different roles and markets.
        """
        self.doc.add_paragraph(modeling_text.strip(), style='CustomBody')

        # Features used
        self.doc.add_paragraph("Features Used", style='CustomHeading2')
        features_text = """
• Location: Geographic cost of living and market demand variations
• Job Title: Role complexity and responsibility level indicators
• Industry: Sector-specific compensation standards and practices
• Experience Years: Career progression and expertise accumulation
• Skills Count: Technical capability breadth and specialization depth
        """
        self.doc.add_paragraph(features_text.strip(), style='CustomBody')

        # Model performance
        self.doc.add_paragraph("Model Performance", style='CustomHeading2')
        r2_score = regression_results.get('test_r2', 0)
        rmse = regression_results.get('test_rmse', 0)

        performance_text = f"""
R² Score: {r2_score:.3f} (explains {r2_score*100:.1f}% of salary variance)
RMSE: ${rmse:,.0f} (average prediction error)
Sample Size: {regression_results.get('sample_size', {}).get('train', 0):,} training records

This performance indicates that our model successfully captures the major factors
influencing salary determination in the job market.
        """
        self.doc.add_paragraph(performance_text.strip(), style='CustomBody')

        # Top salary drivers
        self.doc.add_paragraph("Top Salary Drivers", style='CustomHeading2')

        feature_importance = regression_results.get('feature_importance', pd.DataFrame())
        if not feature_importance.empty:
            drivers_text = "The model identified the following as the most significant salary drivers:\n\n"
            for idx, row in feature_importance.head(5).iterrows():
                drivers_text += f"• {row['feature']}: Coefficient {row['abs_coefficient']:.0f}\n"

            self.doc.add_paragraph(drivers_text.strip(), style='CustomBody')

        # Job seeker implications
        self.doc.add_paragraph("Implications for Job Seekers", style='CustomHeading2')
        implications_text = """
• Identify high-paying locations to target for job searches
• Understand which skills command the highest salary premiums
• Quantify the financial value of experience and specialization
• Compare compensation expectations across different industries and roles
• Make data-driven decisions about career development investments
        """
        self.doc.add_paragraph(implications_text.strip(), style='CustomBody')

    def _add_model_2_section(self, classification_results: Dict[str, Any]):
        """Add Model 2: Classification section."""
        self.doc.add_paragraph("Model 2: Classification for Above-Average Paying Jobs", style='CustomHeading1')

        # What we're modeling
        self.doc.add_paragraph("What We're Modeling", style='CustomHeading2')
        modeling_text = """
We're classifying jobs as "above-average" or "below-average" paying based on location,
title, industry, and skills. This helps identify high-opportunity roles and market segments.
        """
        self.doc.add_paragraph(modeling_text.strip(), style='CustomBody')

        # Model performance
        accuracy = classification_results.get('test_accuracy', 0)
        threshold = classification_results.get('threshold', 0)

        self.doc.add_paragraph("Model Performance", style='CustomHeading2')
        performance_text = f"""
Classification Accuracy: {accuracy:.3f} ({accuracy*100:.1f}% correct predictions)
Above-Average Threshold: ${threshold:,.0f}
Sample Size: {classification_results.get('sample_size', {}).get('train', 0):,} training records

This accuracy level demonstrates the model's effectiveness in identifying
high-opportunity job classifications.
        """
        self.doc.add_paragraph(performance_text.strip(), style='CustomBody')

        # Top predictors
        self.doc.add_paragraph("Top Above-Average Job Predictors", style='CustomHeading2')

        feature_importance = classification_results.get('feature_importance', pd.DataFrame())
        if not feature_importance.empty:
            predictors_text = "The model identified these key predictors of above-average compensation:\n\n"
            for idx, row in feature_importance.head(5).iterrows():
                predictors_text += f"• {row['feature']}: Importance {row['importance']:.3f}\n"

            self.doc.add_paragraph(predictors_text.strip(), style='CustomBody')

        # Job seeker implications
        self.doc.add_paragraph("Implications for Job Seekers", style='CustomHeading2')
        implications_text = """
• Identify which combinations of factors lead to above-average pay
• Target high-opportunity locations and industries for job searches
• Understand which skills unlock premium compensation opportunities
• Focus job search efforts on above-average paying role types
• Develop strategic career plans based on high-value factor combinations
        """
        self.doc.add_paragraph(implications_text.strip(), style='CustomBody')

    def _add_nlp_analysis_section(self, nlp_results: Dict[str, Any]):
        """Add NLP analysis section."""
        self.doc.add_paragraph("Skills Analysis and Natural Language Processing", style='CustomHeading1')

        # Skills extraction
        self.doc.add_paragraph("Skills Extraction and Analysis", style='CustomHeading2')

        insights = nlp_results.get('insights', {})
        skills_text = f"""
Our NLP analysis processed job descriptions and requirements to extract and analyze skills data:

• Total Unique Skills Identified: {insights.get('total_unique_skills', 0):,}
• Top In-Demand Skill: {insights.get('top_skill', 'Not available')}
• Most Valuable Skill (by salary premium): {insights.get('most_valuable_skill', 'Not available')}
• Skill Clusters Created: {insights.get('num_clusters', 0)}

This analysis provides insights into the current skills landscape and identifies
emerging trends in job market requirements.
        """
        self.doc.add_paragraph(skills_text.strip(), style='CustomBody')

        # Skills clustering
        self.doc.add_paragraph("Skills Clustering Results", style='CustomHeading2')

        clusters = nlp_results.get('clusters', {})
        if clusters:
            clustering_text = "Skills were automatically grouped into the following clusters:\n\n"
            for cluster_name, cluster_data in clusters.items():
                clustering_text += f"• {cluster_data.get('description', cluster_name)}: "
                clustering_text += f"{cluster_data.get('total_skills', 0)} skills\n"

            self.doc.add_paragraph(clustering_text.strip(), style='CustomBody')

        # Salary correlation
        correlation_analysis = nlp_results.get('correlation_analysis', pd.DataFrame())
        if not correlation_analysis.empty:
            self.doc.add_paragraph("Skills with Highest Salary Premiums", style='CustomHeading2')

            correlation_text = "Skills analysis revealed the following high-value capabilities:\n\n"
            for _, row in correlation_analysis.head(5).iterrows():
                skill = row.get('skill', 'Unknown')
                premium = row.get('salary_premium', 0)
                correlation_text += f"• {skill}: +${premium:,.0f} salary premium\n"

            self.doc.add_paragraph(correlation_text.strip(), style='CustomBody')

    def _add_strategic_recommendations(self, recommendations: Dict[str, List[str]]):
        """Add strategic recommendations section."""
        self.doc.add_paragraph("Strategic Recommendations", style='CustomHeading1')

        # For job seekers
        if 'for_job_seekers' in recommendations:
            self.doc.add_paragraph("For Job Seekers", style='CustomHeading2')
            job_seeker_text = ""
            for rec in recommendations['for_job_seekers']:
                job_seeker_text += f"• {rec}\n"
            self.doc.add_paragraph(job_seeker_text.strip(), style='CustomBody')

        # For employers
        if 'for_employers' in recommendations:
            self.doc.add_paragraph("For Employers", style='CustomHeading2')
            employer_text = ""
            for rec in recommendations['for_employers']:
                employer_text += f"• {rec}\n"
            self.doc.add_paragraph(employer_text.strip(), style='CustomBody')

        # For policy makers
        if 'for_policy_makers' in recommendations:
            self.doc.add_paragraph("For Policy Makers", style='CustomHeading2')
            policy_text = ""
            for rec in recommendations['for_policy_makers']:
                policy_text += f"• {rec}\n"
            self.doc.add_paragraph(policy_text.strip(), style='CustomBody')

    def _add_technical_appendix(self, analytics_results: Dict[str, Any], nlp_results: Dict[str, Any]):
        """Add technical appendix with detailed methodology."""
        self.doc.add_page_break()
        self.doc.add_paragraph("Technical Appendix", style='CustomHeading1')

        # Methodology
        self.doc.add_paragraph("Methodology", style='CustomHeading2')
        methodology_text = """
This analysis employed machine learning and natural language processing techniques
to analyze job market data and identify salary patterns.

Data Processing:
• Centralized column mapping for consistent data standardization
• Multi-tier data loading strategy (clean → sample → raw)
• Hierarchical missing value imputation
• Geographic data enhancement with coordinate mapping

Machine Learning Models:
• Multiple Linear Regression using scikit-learn
• Random Forest Classification for job categorization
• K-means clustering for skills topic modeling
• TF-IDF vectorization for text analysis

Quality Assurance:
• Cross-validation for model performance assessment
• Feature importance analysis for interpretability
• Error handling with educational value preservation
• Abstraction layer compliance for maintainable architecture
        """
        self.doc.add_paragraph(methodology_text.strip(), style='CustomBody')

        # Data summary
        data_summary = analytics_results.get('data_summary', {})
        if data_summary:
            self.doc.add_paragraph("Data Summary", style='CustomHeading2')

            summary_text = f"""
Total Records Analyzed: {data_summary.get('total_records', 0):,}
Salary Range: ${data_summary.get('salary_range', {}).get('min', 0):,.0f} - ${data_summary.get('salary_range', {}).get('max', 0):,.0f}
Median Salary: ${data_summary.get('salary_range', {}).get('median', 0):,.0f}
Unique Locations: {data_summary.get('unique_locations', 0):,}
Unique Job Titles: {data_summary.get('unique_titles', 0):,}
Unique Industries: {data_summary.get('unique_industries', 0):,}
            """
            self.doc.add_paragraph(summary_text.strip(), style='CustomBody')

        # Limitations and future work
        self.doc.add_paragraph("Limitations and Future Work", style='CustomHeading2')
        limitations_text = """
Limitations:
• Analysis based on job posting data, which may not reflect actual paid salaries
• Geographic analysis focused primarily on US markets
• Skills extraction dependent on job description quality and completeness
• Model performance may vary across different industries and experience levels

Future Work:
• Integration of additional data sources for salary validation
• Expansion to international job markets and currencies
• Real-time model updating with streaming data
• Advanced deep learning models for improved prediction accuracy
• Integration with economic indicators and market trends
        """
        self.doc.add_paragraph(limitations_text.strip(), style='CustomBody')


def generate_comprehensive_docx_report(df: pd.DataFrame = None, output_path: str = "job_market_analytics_report.docx") -> str:
    """
    Convenience function to generate comprehensive DOCX report.

    Args:
        df: Optional DataFrame. If None, uses auto data loading.
        output_path: Path where to save the DOCX report

    Returns:
        Path to generated report
    """
    if not DOCX_AVAILABLE:
        print("[ERROR] python-docx not available. Cannot generate DOCX report.")
        print("Install with: pip install python-docx")
        return ""

    generator = JobMarketReportGenerator(df)
    return generator.generate_complete_report(output_path)
